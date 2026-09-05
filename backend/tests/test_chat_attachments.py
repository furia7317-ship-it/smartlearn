"""Tutor attachments are bounded, extracted and kept separate from persisted chat text."""

from __future__ import annotations

import base64
import io

import pytest
from PIL import Image

from app.services.chat_attachments import AttachmentExtractionError, extract_attachment


def _sample_png_base64() -> str:
    output = io.BytesIO()
    Image.new("RGB", (8, 6), "navy").save(output, format="PNG")
    return base64.b64encode(output.getvalue()).decode("ascii")


def test_text_and_csv_attachments_are_extracted() -> None:
    text = extract_attachment("notes.md", "text/markdown", "# 栈\n后进先出".encode())
    csv_data = extract_attachment("scores.csv", "text/csv", "姓名,成绩\n小林,92".encode())
    assert text["kind"] == "text"
    assert "后进先出" in text["extracted_text"]
    assert "小林 | 92" in csv_data["extracted_text"]


def test_docx_attachment_keeps_paragraphs_and_tables() -> None:
    from docx import Document

    document = Document()
    document.add_heading("二叉树题目", level=1)
    document.add_paragraph("请计算树的高度。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "答案"
    table.cell(0, 1).text = "3"
    buffer = io.BytesIO()
    document.save(buffer)
    result = extract_attachment(
        "problem.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )
    assert result["kind"] == "document"
    assert "请计算树的高度" in result["extracted_text"]
    assert "答案 | 3" in result["extracted_text"]


def test_attachment_rejects_unsupported_or_empty_files() -> None:
    with pytest.raises(AttachmentExtractionError, match="不支持"):
        extract_attachment("legacy.doc", "application/msword", b"binary")
    with pytest.raises(AttachmentExtractionError, match="文件为空"):
        extract_attachment("empty.txt", "text/plain", b"")


def test_native_image_is_attached_to_the_current_user_message() -> None:
    from app.agent.runner import _attach_native_images, _native_image_parts
    from app.schemas.chat import ChatRequest

    request = ChatRequest(
        student_id="student-1",
        message="请解释图里的二叉树",
        attachments=[{
            "id": "attachment-1",
            "name": "tree.png",
            "kind": "image",
            "media_type": "image/png",
            "size": 100,
            "image_data": _sample_png_base64(),
            "recognition_status": "native",
        }],
    )
    image_parts = _native_image_parts(request)
    messages = _attach_native_images(
        [
            {"role": "system", "content": "system"},
            {"role": "user", "content": "请解释图里的二叉树"},
        ],
        image_parts,
    )

    assert image_parts[0]["type"] == "image_url"
    assert image_parts[0]["image_url"]["url"].startswith("data:image/png;base64,")
    assert messages[-1]["content"][0] == {"type": "text", "text": "请解释图里的二叉树"}
    assert messages[-1]["content"][1] == image_parts[0]


def test_native_image_rejects_invalid_or_mismatched_bytes() -> None:
    from app.agent.runner import _native_image_data_url

    with pytest.raises(ValueError, match="base64"):
        _native_image_data_url("not-base64!", "image/png")
    with pytest.raises(ValueError, match="声明格式"):
        _native_image_data_url(_sample_png_base64(), "image/jpeg")


def test_solution_plan_uses_quiz_generator_but_keeps_new_resource_type() -> None:
    from app.schemas.resource import ResourceRequest
    from app.services.planned_resource_pipeline import build_explicit_request_plan

    plan = build_explicit_request_plan(
        ResourceRequest(
            topic="二叉树",
            student_id="student-1",
            material_types=["solution"],
            quiz_config={"choice": 2, "judge": 0, "short": 1},
        ),
        [],
    )
    task = plan.tasks[0]
    assert task.type == "solution"
    assert task.agent == "quiz"
    assert task.quiz_config == {"choice": 2, "judge": 0, "short": 1}
    assert "题目解析" in task.title
